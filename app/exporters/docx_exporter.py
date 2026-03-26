"""
DOCX Exporter Module.

Word belgesi olarak export eder.
Türkçe karakter desteği ve Roadmap2.md formatına uygun.
"""
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

from app.exporters.base_exporter import BaseExporter
from app.schemas.export import (
    ExportSectionEnum, FullReport, SummaryData,
    SEOContentData, AdGroupData, SocialContentData, BrandProfileData
)


class DocxExporter(BaseExporter):
    """
    Word belgesi olarak export eder.
    
    Roadmap2.md formatına uygun:
    - Kapak sayfası
    - İçindekiler
    - Bölüm bazlı içerik
    - Türkçe karakter desteği (UTF-8)
    """
    
    def export(
        self,
        scoring_run_id: int,
        sections: List[ExportSectionEnum],
        filepath: str
    ) -> str:
        """Export işlemini yapar."""
        data = self.collect_data(scoring_run_id, sections)
        
        doc = Document()
        
        # Kapak sayfası
        self._add_cover_page(doc, data)
        
        # İçindekiler
        self._add_toc(doc, sections)
        
        # Bölümler
        if self._should_include(ExportSectionEnum.SUMMARY, sections):
            self._add_summary_section(doc, data.summary)

        if self._should_include(ExportSectionEnum.BRAND_PROFILE, sections):
            self._add_brand_profile_section(doc, data.brand_profile)

        if self._should_include(ExportSectionEnum.SCORING, sections) and data.scoring:
            self._add_scoring_section(doc, data.scoring)
        
        if self._should_include(ExportSectionEnum.CHANNELS, sections) and data.channels:
            self._add_channels_section(doc, data.channels)
        
        if self._should_include(ExportSectionEnum.SEO_CONTENT, sections) and data.seo_contents:
            self._add_seo_content_section(doc, data.seo_contents)
        
        if self._should_include(ExportSectionEnum.ADS, sections) and data.ads:
            self._add_ads_section(doc, data.ads)
        
        if self._should_include(ExportSectionEnum.SOCIAL, sections) and data.social:
            self._add_social_section(doc, data.social)
        
        doc.save(filepath)
        return filepath
    
    def _add_cover_page(self, doc: Document, data: FullReport):
        """Kapak sayfası ekler."""
        # Boşluk
        for _ in range(5):
            doc.add_paragraph()
        
        # Başlık
        title = doc.add_heading('DIGITUS ENGINE RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Alt bilgiler
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f'\nHazırlanma Tarihi: {data.generated_at.strftime("%Y-%m-%d")}\n')
        info.add_run(f'Scoring Run: #{data.scoring_run_id}\n')
        if data.summary:
            info.add_run(f'Toplam Kelime: {data.summary.total_keywords}\n')
        
        doc.add_page_break()
    
    def _add_toc(self, doc: Document, sections: List[ExportSectionEnum]):
        """İçindekiler ekler."""
        doc.add_heading('İÇİNDEKİLER', level=1)
        
        toc_items = []
        if self._should_include(ExportSectionEnum.SUMMARY, sections):
            toc_items.append('1. Ozet')
        if self._should_include(ExportSectionEnum.BRAND_PROFILE, sections):
            toc_items.append('Marka Profili')
        if self._should_include(ExportSectionEnum.SCORING, sections):
            toc_items.append('2. Skorlama Sonuçları')
        if self._should_include(ExportSectionEnum.CHANNELS, sections):
            toc_items.append('3. Kanal Havuzları')
        if self._should_include(ExportSectionEnum.SEO_CONTENT, sections):
            toc_items.append('4. SEO+GEO İçerikleri')
        if self._should_include(ExportSectionEnum.ADS, sections):
            toc_items.append('5. Google Ads Reklam Setleri')
        if self._should_include(ExportSectionEnum.SOCIAL, sections):
            toc_items.append('6. Sosyal Medya İçerikleri')
        
        for item in toc_items:
            doc.add_paragraph(item)
        
        doc.add_page_break()
    
    def _add_summary_section(self, doc: Document, summary: SummaryData):
        """Özet bölümü ekler."""
        doc.add_heading('1. ÖZET', level=1)
        
        doc.add_paragraph(f'Analiz Edilen Kelime: {summary.total_keywords}')
        doc.add_paragraph()
        
        # Kanal dağılımı tablosu
        doc.add_heading('Kanal Dağılımı', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Kanal'
        hdr[1].text = 'Kelime'
        hdr[2].text = 'Stratejik'
        
        channels = [
            ('ADS', summary.ads_count),
            ('SEO', summary.seo_count),
            ('SOCIAL', summary.social_count),
        ]
        
        for channel, count in channels:
            row = table.add_row().cells
            row[0].text = channel
            row[1].text = str(count)
            row[2].text = '-'
        
        doc.add_paragraph()
        
        # Üretilen içerik
        doc.add_heading('Üretilen İçerik', level=2)
        doc.add_paragraph(f'• SEO+GEO Blog Yazısı: {summary.seo_content_count}')
        doc.add_paragraph(f'• Reklam Grubu: {summary.ad_group_count}')
        doc.add_paragraph(f'• Sosyal Medya İçeriği: {summary.social_content_count}')
        
        doc.add_page_break()

    def _add_brand_profile_section(self, doc: Document, brand_profile: Optional[BrandProfileData]):
        """Marka profili bolumunu ekler."""
        doc.add_heading('MARKA PROFILI', level=1)

        if not brand_profile:
            doc.add_paragraph('Bu scoring run icin marka profili bulunamadi.')
            doc.add_page_break()
            return

        doc.add_paragraph(f'Durum: {brand_profile.status}')
        if brand_profile.company_url:
            doc.add_paragraph(f'Firma URL: {brand_profile.company_url}')
        if brand_profile.company_name:
            doc.add_paragraph(f'Firma Adi: {brand_profile.company_name}')
        if brand_profile.sector:
            doc.add_paragraph(f'Sektor: {brand_profile.sector}')
        if brand_profile.target_audience:
            doc.add_paragraph(f'Hedef Kitle: {brand_profile.target_audience}')

        def add_list(title: str, items: list):
            if not items:
                return
            doc.add_heading(title, level=2)
            for item in items:
                doc.add_paragraph(f'- {item}')

        add_list('Urunler', brand_profile.products)
        add_list('Hizmetler', brand_profile.services)
        add_list('Kullanim Alanlari', brand_profile.use_cases)
        add_list('Cozulen Problemler', brand_profile.problems_solved)
        add_list('Marka Terimleri', brand_profile.brand_terms)
        add_list('Dislanacak Temalar', brand_profile.exclude_themes)

        if brand_profile.source_pages:
            doc.add_heading('Kaynak Sayfalar', level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'URL'
            hdr[1].text = 'Baslik'
            hdr[2].text = 'Durum'
            for page in brand_profile.source_pages:
                if not isinstance(page, dict):
                    continue
                row = table.add_row().cells
                row[0].text = str(page.get('url', ''))
                row[1].text = str(page.get('title', ''))
                row[2].text = str(page.get('status', ''))

        if brand_profile.competitors:
            doc.add_heading('Rakip Analizi', level=2)
            comp_table = doc.add_table(rows=1, cols=4)
            comp_table.style = 'Table Grid'
            hdr = comp_table.rows[0].cells
            hdr[0].text = 'URL'
            hdr[1].text = 'Durum'
            hdr[2].text = 'Skor'
            hdr[3].text = 'Ozet'
            for comp in brand_profile.competitors:
                row = comp_table.add_row().cells
                row[0].text = comp.url or '-'
                row[1].text = comp.status or '-'
                row[2].text = f'{comp.consistency_score:.2f}' if comp.consistency_score is not None else '-'
                row[3].text = (comp.summary or '-')[:200]
        elif brand_profile.competitor_urls:
            doc.add_heading('Rakip URLler', level=2)
            for url in brand_profile.competitor_urls:
                doc.add_paragraph(f'- {url}')

        if brand_profile.validation_warnings:
            doc.add_heading('Dogrulama Uyarilari', level=2)
            for warning in brand_profile.validation_warnings:
                doc.add_paragraph(f'- {warning}')

        if brand_profile.error_message:
            doc.add_paragraph(f'Hata Mesaji: {brand_profile.error_message}')

        doc.add_page_break()

    def _add_scoring_section(self, doc: Document, scoring):
        """Skorlama bölümü ekler."""
        doc.add_heading('2. SKORLAMA SONUÇLARI', level=1)
        
        if not scoring.keywords:
            doc.add_paragraph('Skorlama verisi bulunamadı.')
            return
        
        # Tablo
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Kelime'
        hdr[1].text = 'ADS'
        hdr[2].text = 'SEO'
        hdr[3].text = 'SOCIAL'
        hdr[4].text = 'Birincil'
        hdr[5].text = 'Niyet'
        
        for kw in scoring.keywords[:100]:  # İlk 100
            row = table.add_row().cells
            row[0].text = kw.keyword[:30] if kw.keyword else ''
            row[1].text = f'{kw.ads_score:.0f}' if kw.ads_score else '-'
            row[2].text = f'{kw.seo_score:.1f}' if kw.seo_score else '-'
            row[3].text = f'{kw.social_score:.1f}' if kw.social_score else '-'
            row[4].text = kw.primary_channel or '-'
            row[5].text = kw.intent[:15] if kw.intent else '-'
        
        doc.add_page_break()
    
    def _add_channels_section(self, doc: Document, channels):
        """Kanal havuzları bölümü ekler."""
        doc.add_heading('3. KANAL HAVUZLARI', level=1)
        
        for channel_name, pool in [('ADS', channels.ads), ('SEO', channels.seo), ('SOCIAL', channels.social)]:
            doc.add_heading(f'3.{["ADS", "SEO", "SOCIAL"].index(channel_name)+1} {channel_name} Havuzu', level=2)
            
            if not pool.keywords:
                doc.add_paragraph(f'{channel_name} havuzunda kelime bulunmuyor.')
                continue
            
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            hdr = table.rows[0].cells
            hdr[0].text = 'Sıra'
            hdr[1].text = 'Kelime'
            hdr[2].text = 'Skor'
            hdr[3].text = 'Vektör Yakınlığı'
            hdr[4].text = 'Çarpım Skoru'
            hdr[5].text = 'Niyet'
            
            for i, kw in enumerate(pool.keywords[:50], 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = kw.keyword[:40] if kw.keyword else ''
                score = kw.ads_score or kw.seo_score or kw.social_score or 0
                row[2].text = f'{score:.2f}'
                row[3].text = f'{kw.vector_similarity:.3f}' if kw.vector_similarity is not None else '-'
                row[4].text = f'{kw.vector_adjusted_score:.4f}' if kw.vector_adjusted_score is not None else '-'
                row[5].text = kw.intent[:20] if kw.intent else '-'
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_seo_content_section(self, doc: Document, seo_contents):
        """SEO+GEO içerikler bölümü ekler."""
        doc.add_heading('4. SEO+GEO İÇERİKLERİ', level=1)
        
        for c in seo_contents.contents[:20]:  # İlk 20
            doc.add_heading(f'Anahtar Kelime: {c.keyword}', level=2)
            
            doc.add_paragraph(f'BAŞLIK: {c.title}')
            if c.url_suggestion:
                doc.add_paragraph(f'URL ÖNERİSİ: {c.url_suggestion}')
            if c.meta_description:
                doc.add_paragraph(f'META AÇIKLAMA: {c.meta_description}')
            
            # İstatistikler
            stats_parts = []
            if c.word_count:
                stats_parts.append(f'Kelime Sayısı: {c.word_count}')
            if c.keyword_count:
                stats_parts.append(f'Keyword Tekrar: {c.keyword_count}')
            if c.keyword_density:
                stats_parts.append(f'Keyword Yoğunluğu: %{c.keyword_density:.2f}')
            if stats_parts:
                doc.add_paragraph(' | '.join(stats_parts))
            
            # Giriş paragrafı
            if c.intro_paragraph:
                doc.add_heading('GİRİŞ PARAGRAFI', level=3)
                doc.add_paragraph(c.intro_paragraph)
            
            # Body sections (alt başlıklarla birlikte)
            if c.subheadings and c.body_sections:
                doc.add_heading('İÇERİK BÖLÜMLERİ', level=3)
                for i, subheading in enumerate(c.subheadings):
                    p = doc.add_paragraph()
                    run = p.add_run(subheading)
                    run.bold = True
                    if i < len(c.body_sections):
                        doc.add_paragraph(c.body_sections[i])
            elif c.body_content:
                doc.add_heading('İÇERİK', level=3)
                doc.add_paragraph(c.body_content)
            
            # Bullet points
            if c.bullet_points:
                doc.add_heading('MADDELER', level=3)
                for bp in c.bullet_points:
                    text = bp.get('text', '') if isinstance(bp, dict) else str(bp)
                    doc.add_paragraph(f'• {text}')
            
            # Link önerileri
            if c.internal_link_anchor or c.external_link_anchor:
                doc.add_heading('LİNK ÖNERİLERİ', level=3)
                if c.internal_link_anchor:
                    doc.add_paragraph(f'Internal Link: "{c.internal_link_anchor}" → {c.internal_link_url or "-"}')
                if c.external_link_anchor:
                    doc.add_paragraph(f'External Link: "{c.external_link_anchor}" → {c.external_link_url or "-"}')
            
            # Uyumluluk skorları (özet)
            doc.add_heading('UYUMLULUK SKORLARI', level=3)
            score_table = doc.add_table(rows=1, cols=3)
            score_table.style = 'Table Grid'
            
            hdr = score_table.rows[0].cells
            hdr[0].text = 'Metrik'
            hdr[1].text = 'Skor'
            hdr[2].text = 'Durum'
            
            for metric, score in [('SEO', c.seo_score), ('GEO', c.geo_score), ('Combined', c.combined_score)]:
                row = score_table.add_row().cells
                row[0].text = metric
                row[1].text = f'{score:.2f}' if score else '-'
                row[2].text = '✔ Pass' if score and score >= 0.7 else '✗ Fail' if score else '-'
            
            # SEO 11 Kriter Detay Tablosu
            if c.seo_checks:
                doc.add_heading('SEO DETAY (11 Kriter)', level=3)
                seo_detail_table = doc.add_table(rows=1, cols=3)
                seo_detail_table.style = 'Table Grid'
                
                hdr = seo_detail_table.rows[0].cells
                hdr[0].text = '#'
                hdr[1].text = 'Kriter'
                hdr[2].text = 'Durum'
                
                seo_criteria = [
                    ('1', 'Başlık Keyword', 'title_has_keyword'),
                    ('2', 'Başlık Uzunluk (≤70)', 'title_length_ok'),
                    ('3', 'URL Keyword', 'url_has_keyword'),
                    ('4', f'Giriş Keyword (={c.seo_checks.get("intro_keyword_count", "?")})', 'intro_keyword_count'),
                    ('5', 'Kelime Sayısı (300-450)', 'word_count_in_range'),
                    ('6', 'Alt Başlık ≥3', 'subheading_count_ok'),
                    ('7', 'Alt Başlık Keyword', 'subheadings_have_kw'),
                    ('8', 'Internal Link', 'has_internal_link'),
                    ('9', 'External Link', 'has_external_link'),
                    ('10', 'Bullet List', 'has_bullet_list'),
                    ('11', 'Okunabilirlik (≤20 kelime/cümle)', 'sentences_readable'),
                ]
                
                for num, label, key in seo_criteria:
                    row = seo_detail_table.add_row().cells
                    row[0].text = num
                    row[1].text = label
                    val = c.seo_checks.get(key)
                    if key == 'intro_keyword_count':
                        count = val if val is not None else 0
                        row[2].text = f'✔ {count}' if count and count >= 2 else f'✗ {count}'
                    else:
                        row[2].text = '✔ Pass' if val else '✗ Fail'
                
                passed = c.seo_checks.get('total_passed', 0) or 0
                doc.add_paragraph(f'Geçen: {passed}/11')
                if c.seo_checks.get('improvement_notes'):
                    doc.add_paragraph(f'İyileştirme: {c.seo_checks["improvement_notes"]}')
            
            # GEO 7 Kriter Detay Tablosu
            if c.geo_checks:
                doc.add_heading('GEO DETAY (7 Kriter)', level=3)
                geo_detail_table = doc.add_table(rows=1, cols=3)
                geo_detail_table.style = 'Table Grid'
                
                hdr = geo_detail_table.rows[0].cells
                hdr[0].text = '#'
                hdr[1].text = 'Kriter'
                hdr[2].text = 'Durum'
                
                geo_criteria = [
                    ('1', 'Doğrudan Yanıt', 'intro_answers_question'),
                    ('2', 'Snippet Uyumu', 'snippet_extractable'),
                    ('3', 'Bilgi Hiyerarşisi', 'info_hierarchy_strong'),
                    ('4', 'Ton Tutarlılığı', 'tone_is_informative'),
                    ('5', 'Dolgu Yok', 'no_fluff_content'),
                    ('6', 'İlk 50 Kelime Yanıt', 'direct_answer_present'),
                    ('7', 'Doğrulanabilirlik', 'has_verifiable_info'),
                ]
                
                for num, label, key in geo_criteria:
                    row = geo_detail_table.add_row().cells
                    row[0].text = num
                    row[1].text = label
                    val = c.geo_checks.get(key)
                    row[2].text = '✔ Pass' if val else '✗ Fail'
                
                passed = c.geo_checks.get('total_passed', 0) or 0
                doc.add_paragraph(f'Geçen: {passed}/7')
                if c.geo_checks.get('ai_snippet_preview'):
                    doc.add_paragraph(f'AI Snippet: "{c.geo_checks["ai_snippet_preview"][:200]}"')
                if c.geo_checks.get('improvement_notes'):
                    doc.add_paragraph(f'İyileştirme: {c.geo_checks["improvement_notes"]}')
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_ads_section(self, doc: Document, ads):
        """Google Ads bölümü ekler."""
        doc.add_heading('5. GOOGLE ADS REKLAM SETLERİ', level=1)
        
        for g in ads.ad_groups:
            doc.add_heading(f'REKLAM GRUBU: {g.group_name}', level=2)
            
            if g.target_keywords:
                doc.add_paragraph(f'Hedef Kelimeler: {", ".join(g.target_keywords[:5])}')
            
            # Başlıklar
            if g.headlines:
                doc.add_heading(f'BAŞLIKLAR ({len(g.headlines)})', level=3)
                h_table = doc.add_table(rows=1, cols=3)
                h_table.style = 'Table Grid'
                
                hdr = h_table.rows[0].cells
                hdr[0].text = '#'
                hdr[1].text = 'Başlık'
                hdr[2].text = 'Tip'
                
                for i, h in enumerate(g.headlines[:15], 1):
                    row = h_table.add_row().cells
                    row[0].text = str(i)
                    row[1].text = h.headline_text
                    row[2].text = h.headline_type or '-'
            
            # Açıklamalar
            if g.descriptions:
                doc.add_heading(f'AÇIKLAMALAR ({len(g.descriptions)})', level=3)
                for i, d in enumerate(g.descriptions[:4], 1):
                    doc.add_paragraph(f'{i}. {d.description_text}')
            
            # Negatifler
            if g.negative_keywords:
                doc.add_heading(f'NEGATİF KELİMELER ({len(g.negative_keywords)}+)', level=3)
                negatives = ', '.join([n.keyword for n in g.negative_keywords[:10]])
                doc.add_paragraph(negatives)
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_social_section(self, doc: Document, social):
        """Sosyal medya bölümü ekler."""
        doc.add_heading('6. SOSYAL MEDYA İÇERİKLERİ', level=1)
        
        for c in social.contents[:10]:
            doc.add_heading(f'İÇERİK: "{c.idea_title}"', level=2)
            doc.add_paragraph(f'Platform: {c.platform.title()}')
            doc.add_paragraph(f'Viral Potansiyel: {c.trend_alignment:.2f}')
            
            # Hooklar
            if c.hooks:
                doc.add_heading("HOOK'LAR", level=3)
                for i, h in enumerate(c.hooks[:3], 1):
                    doc.add_paragraph(f'{i}. [{h.style}] "{h.text}"')
            
            # Caption
            if c.caption:
                doc.add_heading('CAPTION', level=3)
                doc.add_paragraph(c.caption)
            
            # Hashtagler
            if c.hashtags:
                doc.add_heading("HASHTAG'LER", level=3)
                doc.add_paragraph(' '.join([f'#{t}' for t in c.hashtags[:10]]))
            
            # CTA
            if c.cta_text:
                doc.add_paragraph(f'CTA: "{c.cta_text}"')
            
            doc.add_paragraph()



