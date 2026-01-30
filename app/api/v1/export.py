"""
Export endpoints.
"""
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import tempfile
import os

from app.dependencies import get_db
from app.database.models import ScoringRun, ChannelPool, Keyword, KeywordScore
from app.schemas.export import ExportRequest, ExportFormat, ExportResponse


router = APIRouter()


@router.post("/")
def create_export(
    request: ExportRequest,
    db: Session = Depends(get_db)
):
    """
    Create an export file from scoring run data.
    Supports DOCX, PDF, and Excel formats.
    """
    # Verify scoring run exists
    scoring_run = db.query(ScoringRun).filter(ScoringRun.id == request.scoring_run_id).first()
    if not scoring_run:
        raise HTTPException(status_code=404, detail="Scoring run not found")
    
    # Get data for export
    pools_data = {}
    channels = request.channels or ['ADS', 'SEO', 'SOCIAL']
    
    for channel in channels:
        pools = (
            db.query(ChannelPool, Keyword, KeywordScore)
            .join(Keyword, ChannelPool.keyword_id == Keyword.id)
            .join(
                KeywordScore,
                (KeywordScore.keyword_id == ChannelPool.keyword_id) &
                (KeywordScore.scoring_run_id == ChannelPool.scoring_run_id)
            )
            .filter(ChannelPool.scoring_run_id == request.scoring_run_id)
            .filter(ChannelPool.channel == channel)
            .order_by(ChannelPool.final_rank)
            .all()
        )
        
        pools_data[channel] = [
            {
                'rank': pool.final_rank,
                'keyword': kw.keyword,
                'volume': kw.monthly_volume,
                'sector': kw.sector,
                'score': float(getattr(score, f"{channel.lower()}_score") or 0),
                'is_strategic': pool.is_strategic
            }
            for pool, kw, score in pools
        ]
    
    # Generate file based on format
    if request.format == ExportFormat.EXCEL:
        filepath = _generate_excel(scoring_run, pools_data)
    elif request.format == ExportFormat.DOCX:
        filepath = _generate_docx(scoring_run, pools_data)
    else:  # PDF
        filepath = _generate_pdf(scoring_run, pools_data)
    
    return {
        'message': 'Export created',
        'format': request.format.value,
        'filepath': filepath,
        'channels': list(pools_data.keys()),
        'total_keywords': sum(len(v) for v in pools_data.values())
    }


def _generate_excel(scoring_run: ScoringRun, pools_data: dict) -> str:
    """Generate Excel export."""
    try:
        import openpyxl
        from openpyxl import Workbook
        
        wb = Workbook()
        
        for idx, (channel, data) in enumerate(pools_data.items()):
            if idx == 0:
                ws = wb.active
                ws.title = channel
            else:
                ws = wb.create_sheet(title=channel)
            
            # Headers
            ws.append(['Rank', 'Keyword', 'Volume', 'Sector', 'Score', 'Strategic'])
            
            for item in data:
                ws.append([
                    item['rank'],
                    item['keyword'],
                    item['volume'],
                    item['sector'] or '',
                    item['score'],
                    'Yes' if item['is_strategic'] else ''
                ])
        
        filepath = tempfile.mktemp(suffix='.xlsx')
        wb.save(filepath)
        return filepath
    
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl not installed")


def _generate_docx(scoring_run: ScoringRun, pools_data: dict) -> str:
    """Generate DOCX export."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading(f'DIGITUS ENGINE - Scoring Run #{scoring_run.id}', 0)
        doc.add_paragraph(f'Run Name: {scoring_run.run_name}')
        doc.add_paragraph(f'Status: {scoring_run.status}')
        
        for channel, data in pools_data.items():
            doc.add_heading(f'{channel} Keywords', level=1)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            headers = table.rows[0].cells
            headers[0].text = 'Rank'
            headers[1].text = 'Keyword'
            headers[2].text = 'Volume'
            headers[3].text = 'Score'
            headers[4].text = 'Strategic'
            
            for item in data[:20]:  # Limit to 20 per channel
                row = table.add_row().cells
                row[0].text = str(item['rank'])
                row[1].text = item['keyword']
                row[2].text = str(item['volume'])
                row[3].text = f"{item['score']:.2f}"
                row[4].text = '★' if item['is_strategic'] else ''
        
        filepath = tempfile.mktemp(suffix='.docx')
        doc.save(filepath)
        return filepath
    
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")


def _generate_pdf(scoring_run: ScoringRun, pools_data: dict) -> str:
    """Generate PDF export."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        filepath = tempfile.mktemp(suffix='.pdf')
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        y = height - 50
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, f"DIGITUS ENGINE - Scoring Run #{scoring_run.id}")
        y -= 30
        
        c.setFont("Helvetica", 12)
        for channel, data in pools_data.items():
            if y < 100:
                c.showPage()
                y = height - 50
            
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y, f"{channel} Keywords")
            y -= 20
            
            c.setFont("Helvetica", 10)
            for item in data[:10]:
                c.drawString(70, y, f"{item['rank']}. {item['keyword']} (Score: {item['score']:.2f})")
                y -= 15
            y -= 10
        
        c.save()
        return filepath
    
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab not installed")


@router.get("/download/{filename}")
def download_export(filename: str):
    """Download a generated export file."""
    # Security: validate filename
    if '..' in filename or '/' in filename or '\\' in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    
    filepath = os.path.join(tempfile.gettempdir(), filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(filepath, filename=filename)
